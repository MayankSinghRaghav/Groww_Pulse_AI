"""
Generates a role-specific Word document pulse note as a fallback/alternative to PDF.
"""
import os
import datetime
import logging
from typing import Any, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config.settings import OUTPUT_DIR, APP_NAME

logger = logging.getLogger("word_note")

def generate_word_note(role: str, insights_data: dict, action_ideas: Any, output_path: str):
    """Generate a professional Word document pulse note."""
    try:
        doc = Document()
        
        # Header
        header = doc.sections[0].header
        htable = header.add_table(1, 2, Inches(6.5))
        htab_cells = htable.rows[0].cells
        htab_cells[0].text = f"KUVERA PULSE - {role.upper()}"
        htab_cells[1].text = datetime.datetime.now().strftime("%B %d, %Y")
        htab_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Title
        title = doc.add_heading(f'Weekly Pulse Note: {role}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Intro
        doc.add_paragraph(f"Generated for the {role} to highlight critical feedback themes and action items.")
        
        # Themes
        doc.add_heading('Top 3 Feedback Themes', level=1)
        local_clusters = insights_data.get("local_clusters", {})
        sorted_themes = sorted(local_clusters.items(), key=lambda x: x[1].get("volume", 0), reverse=True)[:3]
        
        for i, (name, data) in enumerate(sorted_themes, 1):
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(f"{name}")
            run.bold = True
            p.add_run(f" ({data.get('volume', 0)} reviews, {data.get('average_rating', 0)}/5 stars)")
            
            # Quotes
            quotes = data.get("representative_quotes", [])
            if quotes:
                doc.add_paragraph(f'"{quotes[0]}"', style='Intense Quote')
        
        # Actions
        doc.add_heading('Recommended Actions', level=1)
        
        # Extract actions correctly
        role_actions = []
        if isinstance(action_ideas, dict):
            role_actions = action_ideas.get(role, [])
        elif isinstance(action_ideas, list):
            # Fallback for old format
            role_actions = action_ideas
        
        # Clean and filter
        import re
        final_actions = []
        for idea in role_actions:
            if not isinstance(idea, str): continue
            cleaned = re.sub(r'^\d+[\.\)]\s*', '', idea).replace("**", "").strip()
            if len(cleaned) > 10 and "action idea" not in cleaned.lower():
                final_actions.append(cleaned)
        
        if not final_actions:
            final_actions = [f"Monitor {role} related signals and prioritize fixes.", "Analyze user feedback trends.", "Draft communication plan."]
            
        for action in final_actions[:3]:
            doc.add_paragraph(action, style='List Bullet')
            
        # Footer
        doc.add_paragraph("\nCONFIDENTIAL | Kuvera Pulse AI Engine")
        
        doc.save(output_path)
        logger.info(f"Word note generated: {output_path}")
        return True
    except Exception as e:
        logger.error(f"Failed to generate Word note: {e}")
        return False
